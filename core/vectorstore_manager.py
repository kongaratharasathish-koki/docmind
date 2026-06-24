# core/vectorstore_manager.py
#
# Handles persistent FAISS vector stores for collections.
# Documents are embedded once and cached on disk under:
#   vectorstores/{user_id}/{collection_id}/faiss_index/
#
# Design notes:
#   - Cache invalidation uses a deterministic hash of document metadata.
#   - If the stored hash matches current documents, the index is loaded from disk.
#   - Otherwise the index is rebuilt and overwritten on disk.
#   - No cloud storage — all indexes are local files.

from __future__ import annotations

from datetime import datetime

import hashlib
import io
from pathlib import Path
from typing import Optional

from core.document_storage import get_collection_path, get_vectorstore_path, list_collection_files
from core.pdf_processor import build_vector_store
from core.session_manager import session_db

import config


def compute_vector_hash(user_id: str, collection_id: int) -> str:
    """
    Compute a deterministic SHA-256 hash from collection document metadata.
    Used to detect when a rebuild is required.
    """
    with session_db._get_connection() as conn:
        cursor = conn.execute(
            "SELECT filename, file_size, content_hash FROM collection_documents WHERE collection_id = ? AND user_id = ?",
            (collection_id, user_id)
        )
        rows = cursor.fetchall()

    raw = "".join(f"{r['filename']}{r['file_size']}{r['content_hash']}" for r in rows)
    return hashlib.sha256(raw.encode()).hexdigest()


def vectorstore_exists(user_id: str, collection_id: int) -> bool:
    """Return True if a saved FAISS index exists for this collection."""
    idx_dir = get_vectorstore_path(user_id, collection_id) / "faiss_index"
    return idx_dir.exists() and (idx_dir / "index.faiss").exists()


# In-process cache: (user_id, collection_id) -> vectorstore
# Bounded by CACHE_MAX_ENTRIES to keep memory predictable. Each upload
# invalidates its own entry; LRU eviction handles cold entries.
_CACHE_MAX_ENTRIES = 64
_vectorstore_cache: "dict[tuple[str, int], object]" = {}
_cache_lru: "list[tuple[str, int]]" = []  # most-recent at the end


def _cache_get(user_id: str, collection_id: int):
    key = (user_id, collection_id)
    vs = _vectorstore_cache.get(key)
    if vs is not None:
        try:
            _cache_lru.remove(key)
        except ValueError:
            pass
        _cache_lru.append(key)
    return vs


def _cache_put(user_id: str, collection_id: int, vs) -> None:
    key = (user_id, collection_id)
    if key in _vectorstore_cache:
        try:
            _cache_lru.remove(key)
        except ValueError:
            pass
    _vectorstore_cache[key] = vs
    _cache_lru.append(key)
    while len(_cache_lru) > _CACHE_MAX_ENTRIES:
        evict = _cache_lru.pop(0)
        _vectorstore_cache.pop(evict, None)


def invalidate_vectorstore_cache(user_id: str, collection_id: int) -> None:
    """Drop a cached vectorstore instance. Call after writes to the collection."""
    key = (user_id, collection_id)
    _vectorstore_cache.pop(key, None)
    try:
        _cache_lru.remove(key)
    except ValueError:
        pass


def load_vectorstore(user_id: str, collection_id: int):
    """
    Load a persisted FAISS index from disk.
    Returns the vector store, or None if it cannot be loaded.
    """
    from langchain_community.vectorstores import FAISS
    from langchain_google_genai import GoogleGenerativeAIEmbeddings

    embeddings = GoogleGenerativeAIEmbeddings(
        model=config.EMBEDDING_MODEL,
        google_api_key=config.GEMINI_API_KEY,
    )

    idx_dir = str(get_vectorstore_path(user_id, collection_id) / "faiss_index")
    try:
        return FAISS.load_local(idx_dir, embeddings, allow_dangerous_deserialization=True)
    except Exception:
        return None


def delete_vectorstore(user_id: str, collection_id: int):
    """Remove the persisted FAISS index for a collection."""
    import shutil
    idx_dir = get_vectorstore_path(user_id, collection_id) / "faiss_index"
    if idx_dir.exists():
        shutil.rmtree(idx_dir)


def _extract_text_from_file(path: Path) -> str:
    """
    Extract plain text from a single .pdf, .txt, or .docx file.
    Returns "" on any extraction failure (caller decides what to do).
    """
    suffix = path.suffix.lower()
    try:
        if suffix == ".pdf":
            from pypdf import PdfReader
            reader = PdfReader(io.BytesIO(path.read_bytes()))
            parts = []
            for page in reader.pages:
                text = page.extract_text() or ""
                if text.strip():
                    parts.append(text.strip())
            return "\n\n".join(parts)

        if suffix == ".txt":
            # Try utf-8 first, fall back to latin-1 (always succeeds for bytes).
            raw = path.read_bytes()
            try:
                return raw.decode("utf-8")
            except UnicodeDecodeError:
                return raw.decode("latin-1", errors="ignore")

        if suffix == ".docx":
            import docx  # python-docx
            document = docx.Document(str(path))
            return "\n\n".join(p.text for p in document.paragraphs if p.text and p.text.strip())

        return ""
    except Exception:
        return ""


def build_collection_vectorstore(user_id: str, collection_id: int):
    """
    Build (or rebuild) a FAISS vector store from the collection's stored documents.
    Accepts .pdf, .txt, and .docx. Per-file errors are isolated: a single
    corrupt file is skipped, the rest still index.
    Returns FAISS vector store and saves it to disk.
    """
    from langchain.schema import Document
    from langchain.text_splitter import RecursiveCharacterTextSplitter
    from langchain_community.vectorstores import FAISS
    from langchain_google_genai import GoogleGenerativeAIEmbeddings

    collection_path = get_collection_path(user_id, collection_id)
    allowed = {".pdf", ".txt", ".docx"}
    files = [
        f for f in collection_path.iterdir()
        if f.is_file() and f.suffix.lower() in allowed
    ]

    if not files:
        raise ValueError("No supported documents found in collection folder.")

    all_pages = []
    skipped = []
    for f in files:
        text = _extract_text_from_file(f)
        if not text or not text.strip():
            skipped.append(f.name)
            continue
        # Plain text / docx is one logical "page 1"; PDFs contribute one
        # entry per page (handled inside _extract_text_from_file).
        if f.suffix.lower() == ".pdf":
            # Re-extract per page so we keep the page numbers for citations.
            try:
                from pypdf import PdfReader
                reader = PdfReader(io.BytesIO(f.read_bytes()))
                for page_num, page in enumerate(reader.pages, start=1):
                    page_text = (page.extract_text() or "").strip()
                    if page_text:
                        all_pages.append({
                            "text": page_text,
                            "source": f.name,
                            "page": page_num,
                        })
            except Exception:
                skipped.append(f.name)
        else:
            all_pages.append({
                "text": text.strip(),
                "source": f.name,
                "page": 1,
            })

    if not all_pages:
        raise ValueError(
            "No extractable text found in collection documents. "
            f"Skipped {len(skipped)} unreadable file(s)."
        )

    splitter = RecursiveCharacterTextSplitter(
        chunk_size=config.CHUNK_SIZE,
        chunk_overlap=config.CHUNK_OVERLAP,
        separators=["\n\n", "\n", ". ", " ", ""],
    )

    documents = []
    for page in all_pages:
        chunks = splitter.split_text(page["text"])
        for chunk in chunks:
            documents.append(Document(
                page_content=chunk,
                metadata={
                    "source": page["source"],
                    "page": page["page"],
                },
            ))

    embeddings = GoogleGenerativeAIEmbeddings(
        model=config.EMBEDDING_MODEL,
        google_api_key=config.GEMINI_API_KEY,
    )

    vector_store = FAISS.from_documents(documents, embeddings)

    # Persist to disk
    idx_dir = get_vectorstore_path(user_id, collection_id) / "faiss_index"
    idx_dir.mkdir(parents=True, exist_ok=True)
    vector_store.save_local(str(idx_dir))

    return vector_store


def load_or_build_collection_vectorstore(user_id: str, collection_id: int):
    """
    Return a FAISS vector store for the collection.
    Load from disk if hash matches; otherwise rebuild and save.
    Also updates collection_metadata.

    A small in-process cache holds the loaded/rebuilt instance keyed by
    (user_id, collection_id) so per-message chat turns don't re-read the
    FAISS index from disk and don't re-construct the embedding client.
    The cache is invalidated by callers (e.g. upload route) whenever the
    underlying documents change.
    """
    from core.collection_manager import ensure_metadata_exists

    ensure_metadata_exists(collection_id)

    current_hash = compute_vector_hash(user_id, collection_id)

    # Fast path: cache hit AND the cache entry was built for the current
    # document set (hash is stored on the instance).
    cached = _cache_get(user_id, collection_id)
    if cached is not None and getattr(cached, "_docmind_hash", None) == current_hash:
        return cached

    # Check if we have a cached index on disk
    if vectorstore_exists(user_id, collection_id):
        with session_db._get_connection() as conn:
            row = conn.execute(
                "SELECT vector_hash FROM collection_metadata WHERE collection_id = ?",
                (collection_id,)
            ).fetchone()
        stored_hash = row["vector_hash"] if row else None

        if stored_hash == current_hash:
            vs = load_vectorstore(user_id, collection_id)
            if vs is not None:
                vs._docmind_hash = current_hash
                _cache_put(user_id, collection_id, vs)
                return vs

    # Rebuild required
    vs = build_collection_vectorstore(user_id, collection_id)
    vs._docmind_hash = current_hash
    _cache_put(user_id, collection_id, vs)

    with session_db._get_connection() as conn:
        conn.execute(
            "UPDATE collection_metadata SET vector_hash = ?, last_indexed_at = ? WHERE collection_id = ?",
            (current_hash, datetime.now().isoformat(), collection_id)
        )
        conn.commit()

    return vs


def build_multi_collection_vectorstore(user_id: str, collection_ids: list):
    """
    Merge multiple collection vector stores into one.
    Used for cross-collection chat where a session references multiple collections.
    """
    from langchain_community.vectorstores import FAISS
    from langchain_google_genai import GoogleGenerativeAIEmbeddings
    import shutil

    embeddings = GoogleGenerativeAIEmbeddings(
        model=config.EMBEDDING_MODEL,
        google_api_key=config.GEMINI_API_KEY,
    )

    # For simplicity, rebuild from disk indexes by merging document sets
    all_documents = []
    for cid in collection_ids:
        idx_dir = get_vectorstore_path(user_id, cid) / "faiss_index"
        if not idx_dir.exists():
            continue
        try:
            vs = FAISS.load_local(str(idx_dir), embeddings, allow_dangerous_deserialization=True)
            docs = vs.similarity_search("", k=1000)
            all_documents.extend(docs)
        except Exception:
            continue

    if not all_documents:
        raise ValueError("No documents found in selected collections.")

    merged = FAISS.from_documents(all_documents, embeddings)
    return merged
